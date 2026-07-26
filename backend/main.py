"""OpenBook AI backend API."""

from __future__ import annotations

import math
import os
import re
import tempfile
from collections.abc import AsyncIterator, Generator
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import Annotated, Callable

import ebooklib
import pymupdf
from bs4 import BeautifulSoup
from docx import Document
from ebooklib import epub
from fastapi import Depends, FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy import Column, Text
from sqlmodel import Field, Session, SQLModel, create_engine, select

BASE_DIRECTORY = Path(__file__).resolve().parent
DATA_DIRECTORY = BASE_DIRECTORY / "data"
DATABASE_PATH = DATA_DIRECTORY / "openbook.db"

MAX_UPLOAD_BYTES = 25 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".epub", ".docx", ".txt"}

DATA_DIRECTORY.mkdir(parents=True, exist_ok=True)

database_url = f"sqlite:///{DATABASE_PATH}"
engine = create_engine(
    database_url,
    connect_args={"check_same_thread": False},
)


def utc_timestamp() -> str:
    """Return the current UTC timestamp."""
    return datetime.now(timezone.utc).isoformat()


class Book(SQLModel, table=True):
    """Stored audiobook source document."""

    id: int | None = Field(default=None, primary_key=True)
    filename: str = Field(index=True)
    file_type: str
    size_bytes: int
    character_count: int
    word_count: int
    estimated_minutes: int
    extracted_text: str = Field(
        sa_column=Column(Text, nullable=False),
    )
    created_at: str = Field(default_factory=utc_timestamp)


class Chapter(SQLModel, table=True):
    """Detected chapter heading belonging to a book."""

    id: int | None = Field(default=None, primary_key=True)
    book_id: int = Field(foreign_key="book.id", index=True)
    position: int
    title: str


@asynccontextmanager
async def lifespan(_: FastAPI) -> AsyncIterator[None]:
    """Create database tables before accepting requests."""
    SQLModel.metadata.create_all(engine)
    yield


app = FastAPI(
    title="OpenBook AI API",
    description="Book processing API for OpenBook AI.",
    version="0.2.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def get_session() -> Generator[Session, None, None]:
    """Provide one database session per request."""
    with Session(engine) as session:
        yield session


DatabaseSession = Annotated[Session, Depends(get_session)]


@app.get("/")
def read_root() -> dict[str, str]:
    """Return API information."""
    return {
        "name": "OpenBook AI API",
        "version": "0.2.0",
    }


@app.get("/health")
def health_check() -> dict[str, str]:
    """Return the backend health status."""
    return {"status": "online"}


@app.get("/books")
def list_books(session: DatabaseSession) -> list[dict[str, object]]:
    """Return all stored books."""
    statement = select(Book).order_by(Book.created_at.desc())
    books = session.exec(statement).all()

    return [
        serialize_book_summary(
            book,
            count_chapters(session, require_book_id(book)),
        )
        for book in books
    ]


@app.get("/books/{book_id}")
def get_book(
    book_id: int,
    session: DatabaseSession,
) -> dict[str, object]:
    """Return one stored book and its detected chapters."""
    book = get_existing_book(session, book_id)
    chapters = get_chapters(session, book_id)

    return serialize_book_detail(book, chapters)


@app.post("/books/upload", status_code=201)
async def upload_book(
    file: Annotated[UploadFile, File(description="Book document")],
    session: DatabaseSession,
) -> dict[str, object]:
    """Extract and permanently store a supported book."""
    filename = Path(file.filename or "uploaded-book").name
    extension = Path(filename).suffix.lower()

    if extension not in ALLOWED_EXTENSIONS:
        supported = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type. Supported types: {supported}",
        )

    temporary_path: Path | None = None

    try:
        temporary_path, size_bytes = await save_upload(file, extension)
        text = normalize_text(extract_text(temporary_path, extension))

        if not text:
            detail = "No readable text was found."

            if extension == ".pdf":
                detail += " Scanned PDFs will require OCR support."

            raise HTTPException(status_code=422, detail=detail)

        words = text.split()
        chapter_titles = detect_chapters(text)

        book = Book(
            filename=filename,
            file_type=extension.removeprefix(".").upper(),
            size_bytes=size_bytes,
            character_count=len(text),
            word_count=len(words),
            estimated_minutes=math.ceil(len(words) / 160),
            extracted_text=text,
        )

        try:
            session.add(book)
            session.commit()
            session.refresh(book)

            book_id = require_book_id(book)

            chapters = [
                Chapter(
                    book_id=book_id,
                    position=index,
                    title=title,
                )
                for index, title in enumerate(
                    chapter_titles,
                    start=1,
                )
            ]

            session.add_all(chapters)
            session.commit()

            for chapter in chapters:
                session.refresh(chapter)

            return serialize_book_detail(book, chapters)
        except Exception:
            session.rollback()
            raise
    finally:
        await file.close()

        if temporary_path is not None:
            temporary_path.unlink(missing_ok=True)


@app.delete("/books/{book_id}")
def delete_book(
    book_id: int,
    session: DatabaseSession,
) -> dict[str, object]:
    """Delete a stored book and its chapter records."""
    book = get_existing_book(session, book_id)
    chapters = get_chapters(session, book_id)

    try:
        for chapter in chapters:
            session.delete(chapter)

        session.delete(book)
        session.commit()
    except Exception:
        session.rollback()
        raise

    return {
        "deleted": True,
        "book_id": book_id,
    }


def get_existing_book(session: Session, book_id: int) -> Book:
    """Return an existing book or raise a 404 response."""
    book = session.get(Book, book_id)

    if book is None:
        raise HTTPException(
            status_code=404,
            detail="Book not found.",
        )

    return book


def require_book_id(book: Book) -> int:
    """Return a persisted book identifier."""
    if book.id is None:
        raise RuntimeError("The book has not been stored.")

    return book.id


def get_chapters(
    session: Session,
    book_id: int,
) -> list[Chapter]:
    """Return chapters in reading order."""
    statement = (
        select(Chapter)
        .where(Chapter.book_id == book_id)
        .order_by(Chapter.position)
    )

    return list(session.exec(statement).all())


def count_chapters(session: Session, book_id: int) -> int:
    """Count detected chapters for a book."""
    return len(get_chapters(session, book_id))


def serialize_book_summary(
    book: Book,
    chapter_count: int,
) -> dict[str, object]:
    """Convert a book into a frontend-safe summary."""
    return {
        "id": require_book_id(book),
        "filename": book.filename,
        "file_type": book.file_type,
        "size_bytes": book.size_bytes,
        "character_count": book.character_count,
        "word_count": book.word_count,
        "estimated_minutes": book.estimated_minutes,
        "chapter_count": chapter_count,
        "created_at": book.created_at,
    }


def serialize_book_detail(
    book: Book,
    chapters: list[Chapter],
) -> dict[str, object]:
    """Convert a book and chapters into a detailed response."""
    response = serialize_book_summary(book, len(chapters))

    response.update(
        {
            "preview": book.extracted_text[:5000],
            "chapters": [
                {
                    "id": chapter.id,
                    "position": chapter.position,
                    "title": chapter.title,
                }
                for chapter in chapters
            ],
        }
    )

    return response


async def save_upload(
    upload: UploadFile,
    extension: str,
) -> tuple[Path, int]:
    """Store an upload temporarily while enforcing its size limit."""
    descriptor, filename = tempfile.mkstemp(suffix=extension)
    os.close(descriptor)

    path = Path(filename)
    size_bytes = 0

    try:
        with path.open("wb") as destination:
            while chunk := await upload.read(1024 * 1024):
                size_bytes += len(chunk)

                if size_bytes > MAX_UPLOAD_BYTES:
                    raise HTTPException(
                        status_code=413,
                        detail="The maximum upload size is 25 MB.",
                    )

                destination.write(chunk)

        return path, size_bytes
    except Exception:
        path.unlink(missing_ok=True)
        raise


def extract_text(path: Path, extension: str) -> str:
    """Extract text using the correct document reader."""
    extractors: dict[str, Callable[[Path], str]] = {
        ".txt": extract_txt,
        ".pdf": extract_pdf,
        ".docx": extract_docx,
        ".epub": extract_epub,
    }

    try:
        return extractors[extension](path)
    except HTTPException:
        raise
    except Exception as error:
        raise HTTPException(
            status_code=422,
            detail=f"Could not process the document: {error}",
        ) from error


def extract_txt(path: Path) -> str:
    """Read a plain-text file."""
    content = path.read_bytes()

    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise ValueError("The text encoding could not be recognized.")


def extract_pdf(path: Path) -> str:
    """Read text from a text-based PDF."""
    with pymupdf.open(path) as document:
        return "\n\n".join(
            page.get_text("text")
            for page in document
        )


def extract_docx(path: Path) -> str:
    """Read paragraphs and tables from a DOCX file."""
    document = Document(path)
    sections: list[str] = []

    sections.extend(
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    )

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )

            if row_text:
                sections.append(row_text)

    return "\n\n".join(sections)


def extract_epub(path: Path) -> str:
    """Read text from EPUB document sections."""
    book = epub.read_epub(str(path))
    sections: list[str] = []

    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "html.parser")

        for element in soup(["script", "style", "nav"]):
            element.decompose()

        section = soup.get_text("\n", strip=True)

        if section:
            sections.append(section)

    return "\n\n".join(sections)


def normalize_text(text: str) -> str:
    """Normalize whitespace while retaining paragraphs."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def detect_chapters(text: str) -> list[str]:
    """Detect common chapter headings."""
    pattern = re.compile(
        r"^(chapter|part|book|section)\s+"
        r"([0-9]+|[ivxlcdm]+|one|two|three|four|five|six|seven|eight|nine|ten)"
        r"(?:\s*[:.-]\s*|\s*$|\s+.+)",
        re.IGNORECASE,
    )

    chapters: list[str] = []
    seen: set[str] = set()

    for line in text.splitlines():
        candidate = line.strip()

        if not candidate or len(candidate) > 120:
            continue

        if pattern.match(candidate):
            key = candidate.casefold()

            if key not in seen:
                chapters.append(candidate)
                seen.add(key)

    return chapters
