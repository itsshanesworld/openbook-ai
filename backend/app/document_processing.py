"""Document extraction and narration splitting services."""

from __future__ import annotations

import os
import posixpath
import re
import tempfile
from pathlib import Path
from urllib.parse import unquote
from typing import Callable

import ebooklib
import pymupdf
from bs4 import BeautifulSoup
from docx import Document
from ebooklib import epub
from fastapi import HTTPException, UploadFile

MAX_UPLOAD_BYTES = 25 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".epub", ".docx", ".txt"}


async def save_upload(
    upload: UploadFile,
    extension: str,
) -> tuple[Path, int]:
    """Store an upload temporarily while enforcing its size limit."""
    descriptor, filename = tempfile.mkstemp(suffix=extension)
    os.close(descriptor)

    path = Path(filename)
    size_bytes = 0

    try:
        with path.open("wb") as destination:
            while chunk := await upload.read(1024 * 1024):
                size_bytes += len(chunk)

                if size_bytes > MAX_UPLOAD_BYTES:
                    raise HTTPException(
                        status_code=413,
                        detail="The maximum upload size is 25 MB.",
                    )

                destination.write(chunk)

        return path, size_bytes
    except Exception:
        path.unlink(missing_ok=True)
        raise


def extract_text(path: Path, extension: str) -> str:
    """Extract text using the appropriate document reader."""
    extractors: dict[str, Callable[[Path], str]] = {
        ".txt": extract_txt,
        ".pdf": extract_pdf,
        ".docx": extract_docx,
        ".epub": extract_epub,
    }

    try:
        return extractors[extension](path)
    except HTTPException:
        raise
    except Exception as error:
        raise HTTPException(
            status_code=422,
            detail=f"Could not process the document: {error}",
        ) from error


def extract_txt(path: Path) -> str:
    """Read a plain-text document."""
    content = path.read_bytes()

    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise ValueError("The text encoding could not be recognized.")


def extract_pdf(path: Path) -> str:
    """Read PDF text and preserve outline chapter boundaries."""
    with pymupdf.open(path) as document:
        toc_entries = get_pdf_toc_entries(
            document
        )

        sections: list[str] = []

        for page_index, page in enumerate(
            document,
            start=1,
        ):
            page_text = page.get_text(
                "text"
            ).strip()

            existing_lines = {
                line.strip().casefold()
                for line in page_text.splitlines()
                if line.strip()
            }

            page_titles = [
                title
                for title, page_number in toc_entries
                if page_number == page_index
                and title.casefold()
                not in existing_lines
            ]

            page_parts = [
                *page_titles,
            ]

            if page_text:
                page_parts.append(
                    page_text
                )

            section = "\n".join(
                page_parts
            ).strip()

            if section:
                sections.append(
                    section
                )

    return "\n\n".join(
        sections
    )


def extract_pdf_chapter_titles(
    path: Path,
) -> list[str]:
    """Extract chapter titles from the PDF outline."""
    with pymupdf.open(path) as document:
        return [
            title
            for title, _ in get_pdf_toc_entries(
                document
            )
        ]


def get_pdf_toc_entries(
    document: pymupdf.Document,
) -> list[tuple[str, int]]:
    """Return usable PDF outline titles and target pages."""
    entries: list[tuple[str, int]] = []
    seen: set[tuple[str, int]] = set()

    for entry in document.get_toc(
        simple=True
    ):
        if len(entry) < 3:
            continue

        title = clean_document_metadata_value(
            entry[1]
        )

        page_number = entry[2]

        if title is None:
            continue

        if not isinstance(
            page_number,
            int,
        ):
            continue

        if (
            page_number < 1
            or page_number > document.page_count
        ):
            continue

        key = (
            title.casefold(),
            page_number,
        )

        if key in seen:
            continue

        seen.add(
            key
        )

        entries.append(
            (
                title,
                page_number,
            )
        )

    return entries


def extract_pdf_metadata(
    path: Path,
) -> tuple[str | None, str | None]:
    """Extract structured title and author metadata from PDF."""
    with pymupdf.open(path) as document:
        metadata = document.metadata or {}

    title = clean_document_metadata_value(
        metadata.get("title")
    )

    author = clean_document_metadata_value(
        metadata.get("author")
    )

    return title, author


def extract_pdf_cover(
    path: Path,
) -> tuple[str, bytes] | None:
    """Render the first PDF page as PNG cover artwork."""
    with pymupdf.open(path) as document:
        if document.page_count < 1:
            return None

        page = document[0]

        pixmap = page.get_pixmap(
            dpi=144,
            alpha=False,
        )

        content = pixmap.tobytes("png")

    if not content.startswith(
        b"\x89PNG\r\n\x1a\n"
    ):
        return None

    return "cover.png", content


def extract_docx(path: Path) -> str:
    """Read paragraphs and tables from a DOCX file."""
    document = Document(path)
    sections: list[str] = []

    sections.extend(
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    )

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )

            if row_text:
                sections.append(row_text)

    return "\n\n".join(sections)


def extract_docx_chapter_titles(
    path: Path,
) -> list[str]:
    """Extract chapter titles from Word Heading styles."""
    document = Document(path)
    chapters: list[str] = []
    seen: set[str] = set()

    heading_pattern = re.compile(
        r"^Heading [1-9]$",
        re.IGNORECASE,
    )

    for paragraph in document.paragraphs:
        title = clean_document_metadata_value(
            paragraph.text
        )

        if title is None:
            continue

        style = paragraph.style

        style_name = (
            style.name
            if style is not None
            else ""
        )

        if not heading_pattern.fullmatch(
            style_name
        ):
            continue

        key = title.casefold()

        if key in seen:
            continue

        seen.add(
            key
        )

        chapters.append(
            title
        )

    return chapters


def extract_docx_metadata(
    path: Path,
) -> tuple[str | None, str | None]:
    """Extract structured title and author metadata from DOCX."""
    document = Document(path)
    properties = document.core_properties

    title = clean_document_metadata_value(
        properties.title
    )

    author = clean_document_metadata_value(
        properties.author
    )

    if (
        author is not None
        and author.casefold() == "python-docx"
    ):
        author = None

    return title, author


def extract_epub(path: Path) -> str:
    """Read EPUB text in spine order with official TOC headings."""
    book = epub.read_epub(str(path))
    toc_entries = get_epub_toc_entries(book)
    sections: list[str] = []

    for item in get_epub_reading_order(book):
        soup = BeautifulSoup(
            item.get_content(),
            "html.parser",
        )

        for element in soup(
            ["script", "style", "nav"]
        ):
            element.decompose()

        item_name = normalize_epub_document_name(
            get_epub_item_name(item)
        )

        matching_entries = [
            (
                title,
                fragment,
            )
            for title, href, fragment in toc_entries
            if epub_document_names_match(
                item_name,
                normalize_epub_document_name(href),
            )
        ]

        inject_epub_toc_titles(
            soup,
            matching_entries,
        )

        section = soup.get_text(
            "\n",
            strip=True,
        )

        if section:
            sections.append(section)

    return "\n\n".join(sections)


def extract_epub_toc_titles(
    path: Path,
) -> list[str]:
    """Return official playable chapter titles from an EPUB TOC."""
    book = epub.read_epub(str(path))

    return [
        title
        for title, _, _ in get_epub_toc_entries(book)
    ]


def get_epub_toc_entries(
    book: epub.EpubBook,
) -> list[tuple[str, str, str | None]]:
    """Flatten usable EPUB TOC links while preserving TOC order."""
    entries: list[tuple[str, str, str | None]] = []
    seen: set[tuple[str, str, str | None]] = set()

    def add_node(node: object) -> None:
        title = clean_document_metadata_value(
            getattr(
                node,
                "title",
                None,
            )
        )

        href_value = getattr(
            node,
            "href",
            None,
        )

        if href_value is None:
            name_getter = getattr(
                node,
                "get_name",
                None,
            )

            if callable(name_getter):
                href_value = name_getter()

        href = clean_document_metadata_value(
            href_value
        )

        if title is None or href is None:
            return

        document_href, fragment = split_epub_href(
            href
        )

        if not document_href:
            return

        key = (
            title.casefold(),
            normalize_epub_document_name(
                document_href
            ),
            fragment.casefold()
            if fragment
            else None,
        )

        if key in seen:
            return

        seen.add(key)

        entries.append(
            (
                title,
                document_href,
                fragment,
            )
        )

    def walk(nodes: object) -> None:
        if not isinstance(
            nodes,
            (list, tuple),
        ):
            add_node(nodes)
            return

        for node in nodes:
            if (
                isinstance(node, tuple)
                and len(node) == 2
                and isinstance(
                    node[1],
                    (list, tuple),
                )
            ):
                parent, children = node

                add_node(parent)
                walk(children)
                continue

            if isinstance(node, list):
                walk(node)
                continue

            add_node(node)

    walk(book.toc)

    return entries


def get_epub_reading_order(
    book: epub.EpubBook,
) -> list[object]:
    """Return readable EPUB documents using spine order first."""
    documents: list[object] = []
    seen_names: set[str] = set()

    def add_item(item: object | None) -> None:
        if item is None:
            return

        type_getter = getattr(
            item,
            "get_type",
            None,
        )

        if (
            not callable(type_getter)
            or type_getter()
            != ebooklib.ITEM_DOCUMENT
        ):
            return

        if isinstance(
            item,
            epub.EpubNav,
        ):
            return

        name = normalize_epub_document_name(
            get_epub_item_name(item)
        )

        if not name or name in seen_names:
            return

        seen_names.add(name)
        documents.append(item)

    for spine_entry in book.spine:
        item_reference = (
            spine_entry[0]
            if isinstance(
                spine_entry,
                (list, tuple),
            )
            and spine_entry
            else spine_entry
        )

        if hasattr(
            item_reference,
            "get_content",
        ):
            item = item_reference
        else:
            item = book.get_item_with_id(
                str(item_reference)
            )

        add_item(item)

    for item in book.get_items_of_type(
        ebooklib.ITEM_DOCUMENT
    ):
        add_item(item)

    return documents


def get_epub_item_name(
    item: object,
) -> str:
    """Return an EPUB manifest item's document name."""
    name_getter = getattr(
        item,
        "get_name",
        None,
    )

    if callable(name_getter):
        return str(
            name_getter()
        )

    return str(
        getattr(
            item,
            "file_name",
            "",
        )
    )


def split_epub_href(
    href: str,
) -> tuple[str, str | None]:
    """Split an EPUB href into document path and fragment."""
    document_href, separator, fragment = (
        href.partition("#")
    )

    decoded_document = unquote(
        document_href
    )

    decoded_fragment = (
        unquote(fragment)
        if separator and fragment
        else None
    )

    return (
        decoded_document,
        decoded_fragment,
    )


def normalize_epub_document_name(
    value: str,
) -> str:
    """Normalize an EPUB document path for matching."""
    document_name = unquote(
        value.split(
            "#",
            1,
        )[0]
    ).replace(
        "\\",
        "/",
    )

    normalized = posixpath.normpath(
        document_name
    )

    if normalized == ".":
        return ""

    while normalized.startswith(
        "../"
    ):
        normalized = normalized[3:]

    return normalized.lstrip("/")


def epub_document_names_match(
    item_name: str,
    toc_name: str,
) -> bool:
    """Match equivalent manifest and TOC document paths."""
    if not item_name or not toc_name:
        return False

    if item_name == toc_name:
        return True

    return (
        item_name.endswith(
            f"/{toc_name}"
        )
        or toc_name.endswith(
            f"/{item_name}"
        )
    )


def inject_epub_toc_titles(
    soup: BeautifulSoup,
    entries: list[tuple[str, str | None]],
) -> None:
    """Ensure official TOC titles appear at chapter boundaries."""
    existing_lines = {
        line.strip().casefold()
        for line in soup.get_text(
            "\n",
            strip=True,
        ).splitlines()
        if line.strip()
    }

    for title, fragment in entries:
        if title.casefold() in existing_lines:
            continue

        marker = soup.new_tag("p")
        marker["data-openbook-toc-title"] = "true"
        marker.string = title

        target = None

        if fragment:
            target = soup.find(
                id=fragment
            )

            if target is None:
                target = soup.find(
                    attrs={
                        "name": fragment,
                    }
                )

        if target is not None:
            target.insert_before(
                marker
            )
        else:
            container = (
                soup.body
                if soup.body is not None
                else soup
            )

            container.insert(
                0,
                marker,
            )

        existing_lines.add(
            title.casefold()
        )


def extract_epub_metadata(
    path: Path,
) -> tuple[str | None, str | None]:
    """Extract structured title and author metadata from EPUB."""
    book = epub.read_epub(
        str(path)
    )

    title = first_epub_metadata_value(
        book.get_metadata(
            "DC",
            "title",
        )
    )

    authors: list[str] = []
    seen_authors: set[str] = set()

    for value, _ in book.get_metadata(
        "DC",
        "creator",
    ):
        author = clean_document_metadata_value(
            value
        )

        if author is None:
            continue

        key = author.casefold()

        if key in seen_authors:
            continue

        authors.append(author)
        seen_authors.add(key)

    author = (
        ", ".join(authors)
        if authors
        else None
    )

    return title, author


def first_epub_metadata_value(
    values: list[tuple[object, dict[str, str]]],
) -> str | None:
    """Return the first usable EPUB metadata value."""
    for value, _ in values:
        cleaned = clean_document_metadata_value(
            value
        )

        if cleaned is not None:
            return cleaned

    return None


def clean_document_metadata_value(
    value: object,
) -> str | None:
    """Normalize one structured document metadata value."""
    if value is None:
        return None

    cleaned = re.sub(
        r"\s+",
        " ",
        str(value),
    ).strip()

    if not cleaned:
        return None

    return cleaned


def extract_epub_cover(
    path: Path,
) -> tuple[str, bytes] | None:
    """Extract JPG or PNG cover artwork from an EPUB."""
    book = epub.read_epub(str(path))

    cover_item = find_epub_cover_item(book)

    if cover_item is None:
        return None

    content = bytes(
        cover_item.get_content()
    )

    extension = detect_epub_cover_extension(
        content
    )

    if extension is None:
        return None

    name_getter = getattr(
        cover_item,
        "get_name",
        None,
    )

    if callable(name_getter):
        item_name = str(name_getter())
    else:
        item_name = str(
            getattr(
                cover_item,
                "file_name",
                "cover",
            )
        )

    filename = Path(item_name).name

    if not filename.casefold().endswith(
        (
            ".jpg",
            ".jpeg",
            ".png",
        )
    ):
        filename = f"cover{extension}"

    return filename, content


def find_epub_cover_item(
    book: epub.EpubBook,
) -> object | None:
    """Find the most likely EPUB cover manifest item."""
    candidates: list[object] = []
    seen_ids: set[int] = set()

    def add_candidate(
        item: object | None,
    ) -> None:
        if item is None:
            return

        identity = id(item)

        if identity in seen_ids:
            return

        candidates.append(item)
        seen_ids.add(identity)

    for item in book.get_items_of_type(
        ebooklib.ITEM_COVER
    ):
        add_candidate(item)

    for _, attributes in book.get_metadata(
        "OPF",
        "cover",
    ):
        cover_id = attributes.get("content")

        if cover_id:
            add_candidate(
                book.get_item_with_id(
                    cover_id
                )
            )

    all_items = list(book.get_items())

    for item in all_items:
        properties = getattr(
            item,
            "properties",
            [],
        )

        if isinstance(properties, str):
            property_names = properties.split()
        else:
            property_names = list(
                properties or []
            )

        if "cover-image" in property_names:
            add_candidate(item)

    for item in book.get_items_of_type(
        ebooklib.ITEM_IMAGE
    ):
        name_getter = getattr(
            item,
            "get_name",
            None,
        )

        if callable(name_getter):
            item_name = str(
                name_getter()
            )
        else:
            item_name = str(
                getattr(
                    item,
                    "file_name",
                    "",
                )
            )

        if "cover" in item_name.casefold():
            add_candidate(item)

    for candidate in candidates:
        try:
            content = bytes(
                candidate.get_content()
            )
        except Exception:
            continue

        if (
            detect_epub_cover_extension(
                content
            )
            is not None
        ):
            return candidate

    return None


def detect_epub_cover_extension(
    content: bytes,
) -> str | None:
    """Recognize cover formats supported by OpenBook AI."""
    if content.startswith(
        b"\x89PNG\r\n\x1a\n"
    ):
        return ".png"

    if content.startswith(
        b"\xff\xd8\xff"
    ):
        return ".jpg"

    return None


def normalize_text(text: str) -> str:
    """Normalize whitespace while retaining paragraphs."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def detect_chapters(text: str) -> list[str]:
    """Detect common numbered and standalone book headings."""
    numbered_pattern = re.compile(
        r"^(chapter|part|book|section)\s+"
        r"([0-9]+|[ivxlcdm]+|one|two|three|four|five|six|seven|eight|nine|ten)"
        r"(?:\s*[:.-]\s*|\s*$|\s+.+)",
        re.IGNORECASE,
    )

    standalone_pattern = re.compile(
        r"^(introduction|prologue|preface|foreword|epilogue|afterword)"
        r"(?:\s*[:.-]\s*.+)?$",
        re.IGNORECASE,
    )

    chapters: list[str] = []
    seen: set[str] = set()

    for line in text.splitlines():
        candidate = line.strip()

        if not candidate or len(candidate) > 120:
            continue

        is_heading = (
            numbered_pattern.match(candidate)
            or standalone_pattern.match(candidate)
        )

        if not is_heading:
            continue

        key = candidate.casefold()

        if key not in seen:
            chapters.append(candidate)
            seen.add(key)

    return chapters


def split_into_narration_sections(
    text: str,
    target_words: int = 350,
    chapter_headings: list[str] | None = None,
) -> list[str]:
    """Split narration while keeping chapter boundaries intact."""
    normalized_text = normalize_text(text)

    if not normalized_text:
        return []

    detected_headings = detect_chapters(
        normalized_text
    )

    all_chapter_headings = {
        heading.casefold()
        for heading in detected_headings
    }

    if chapter_headings:
        all_chapter_headings.update(
            heading.casefold()
            for heading in chapter_headings
            if heading.strip()
        )

    paragraphs = [
        paragraph.strip()
        for paragraph in normalized_text.split("\n\n")
        if paragraph.strip()
    ]

    sections: list[str] = []
    current_parts: list[str] = []
    current_word_count = 0
    chapter_heading_seen = False

    def flush_current_section() -> None:
        nonlocal current_parts
        nonlocal current_word_count

        if current_parts:
            sections.append("\n\n".join(current_parts).strip())
            current_parts = []
            current_word_count = 0

    def add_piece(piece: str) -> None:
        nonlocal current_word_count

        piece = piece.strip()

        if not piece:
            return

        piece_word_count = len(piece.split())

        if (
            current_parts
            and current_word_count + piece_word_count > target_words
        ):
            flush_current_section()

        current_parts.append(piece)
        current_word_count += piece_word_count

        if current_word_count >= target_words:
            flush_current_section()

    def add_size_aware_piece(piece: str) -> None:
        piece = piece.strip()

        if not piece:
            return

        if len(piece.split()) <= target_words:
            add_piece(piece)
            return

        sentences = [
            sentence.strip()
            for sentence in re.split(
                r"(?<=[.!?])\s+",
                piece,
            )
            if sentence.strip()
        ]

        for sentence in sentences:
            words = sentence.split()

            if len(words) <= target_words:
                add_piece(sentence)
                continue

            for start in range(0, len(words), target_words):
                add_piece(
                    " ".join(words[start : start + target_words])
                )

    def split_at_chapter_headings(
        paragraph: str,
    ) -> list[tuple[str, bool]]:
        lines = [
            line.strip()
            for line in paragraph.splitlines()
            if line.strip()
        ]

        if not lines:
            return []

        pieces: list[tuple[str, bool]] = []
        current_lines: list[str] = []
        starts_with_heading = False

        for line in lines:
            is_heading = (
                line.casefold()
                in all_chapter_headings
            )

            if is_heading:
                if current_lines:
                    pieces.append(
                        (
                            "\n".join(current_lines),
                            starts_with_heading,
                        )
                    )

                current_lines = [line]
                starts_with_heading = True
                continue

            current_lines.append(line)

        if current_lines:
            pieces.append(
                (
                    "\n".join(current_lines),
                    starts_with_heading,
                )
            )

        return pieces

    for paragraph in paragraphs:
        pieces = split_at_chapter_headings(paragraph)

        for piece, starts_with_heading in pieces:
            if starts_with_heading:
                if chapter_heading_seen:
                    flush_current_section()

                chapter_heading_seen = True

            add_size_aware_piece(piece)

    flush_current_section()

    return sections
